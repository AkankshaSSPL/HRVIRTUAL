from docx import Document
from docx.shared import Pt, Inches
import os

doc = Document()
doc.add_heading('VirtualHR System Overview', 0)

# Intro & Dual Operation Mode
p = doc.add_paragraph()
p.add_run('The VirtualHR platform allows HR administrators to create, update, and comprehensively view employee profiles. A key feature of the system is its Dual Operation Mode:').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run('Manual UI Operation: ').bold = True
p.add_run('Users can navigate traditional forms, drawers, and dynamic dashboards to perform all tasks manually. The UI provides real-time format validation to prevent erroneous data entry.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('AI Agent Command: ').bold = True
p.add_run("Users can simply type natural language commands to instantly execute complex workflows. Examples include 'Add 5 days of sick leave for John', 'Update Sarah\\'s bank account to 123456789', or 'Generate payroll for August'. This keeps operations extremely fast, reduces clicks, and minimizes token count usage by streamlining actions.")

p = doc.add_paragraph(style='List Bullet')
p.add_run('Real-Time Drawer Updates: ').bold = True
p.add_run('Profiles can be quickly updated through a slide-out drawer interface without losing the context of the main dashboard. Partial updates are supported instantly.')

doc.add_heading('1. Employee Profile & Onboarding', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('HR can easily create new profiles and view comprehensive details including personal data, employment history, and statutory banking information.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('The platform tracks new hires as they submit required details across Personal, Employment, Documents, and Banking tabs.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('A visual onboarding status panel shows exactly what percentage of the profile is complete and lists the exact missing fields.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Document Uploads: Features secure digital uploads for identity documents like PAN and Aadhaar, which are tagged with verification statuses.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('The system prevents employees from becoming fully active until all compliance checks are met, enforcing strict data governance.')

doc.add_heading('2. Seating and Assets Allocation', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run("HR can assign physical office seats directly from an employee's profile, ensuring accurate facility management.")
p = doc.add_paragraph(style='List Bullet')
p.add_run('Equipment like laptops, monitors, keyboards, and access cards are digitally mapped to the employee using unique Asset IDs.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Tracks the exact validity dates and current condition status of every assigned asset.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('A dedicated asset modal makes it easy to issue new equipment or mark items as returned.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Crucially prevents former employees from retaining company property after offboarding by linking clearance to asset returns.')

doc.add_heading('3. Payroll Generation (Employee & Consultant)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Automates monthly salary calculations uniquely tailored based on the worker's employment contract type.")
p = doc.add_paragraph(style='List Bullet')
p.add_run('Employees: ').bold = True
p.add_run('Processes complex fixed pay components (Basic, HRA, Allowances) and applies statutory deductions like TDS and Provident Fund (PF).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Consultants: ').bold = True
p.add_run('Calculates straightforward professional fees or stipends without applying traditional employee tax deductions.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('TDS Configuration: Allows HR to assign specific Tax Deducted at Source brackets (New vs Old regime) for individualized tax compliance.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Provides a detailed preview breakdown of earnings and deductions before finalizing the monthly ledger to catch errors early.')

doc.add_heading('4. Attendance and Leave Management', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Captures daily employee check-ins and check-outs in real-time, logging precise hours worked.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Provides a monthly calendar view showing days marked as Present, Absent, or Weekend.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Manages requested time off, multi-level approval workflows, and automatically deducts approved days from annual leave balances.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Supports different leave types (Sick, Casual, Earned) with specific accrual rules.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Gives supervisors an up-to-date view of who is present without needing to cross-reference manual spreadsheets.')

doc.add_heading('5. Payroll Integration with Attendance', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('The HR operational system is tightly connected directly to the financial payroll engine.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('At the end of the month, payroll automatically pulls the finalized attendance and leave data without any manual exports.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Calculates Loss of Pay (LOP) for any unapproved absences instantly, deducting it from the gross pay.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Handles mid-month joiners seamlessly by pro-rating their first salary automatically based on their joining date.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Removes the need for HR and Finance to manually match records, eliminating human error in salary payouts.')

doc.add_paragraph('\n[ Insert Screenshots Here ]')

doc_path = os.path.join(os.path.expanduser("~"), "Desktop", "VirtualHR_System_Overview_Expanded.docx")
doc.save(doc_path)
print(f"Document generated successfully at {doc_path}")
